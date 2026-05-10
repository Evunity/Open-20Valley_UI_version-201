"""Generate the four Open Valley handover Word (.docx) files.

Run from the repo root:
    py scripts/generate_handover_docx.py
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor


# ---------- Theme ----------

PRIMARY = RGBColor(0x1F, 0x4E, 0x79)       # deep blue
ACCENT = RGBColor(0x2E, 0x86, 0xAB)        # teal
TEXT = RGBColor(0x1A, 0x1A, 0x1A)
MUTED = RGBColor(0x55, 0x55, 0x55)
TABLE_HEADER_BG = "1F4E79"
ZEBRA_BG = "F3F6FB"
CODE_BG = "F5F5F7"
CALLOUT_NOTE_BG = "EAF3FB"
CALLOUT_NOTE_BAR = "1F4E79"
CALLOUT_IMPORTANT_BG = "FFF4CE"
CALLOUT_IMPORTANT_BAR = "B58105"
CALLOUT_TIP_BG = "EAF7EE"
CALLOUT_TIP_BAR = "2E7D32"


# ---------- Low-level helpers ----------


def _shade(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _set_cell_borders(cell, color: str = "DDDDDD", size: int = 4) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), str(size))
        b.set(qn("w:color"), color)
        borders.append(b)
    tc_pr.append(borders)


def _paragraph_shade(paragraph, hex_color: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    p_pr.append(shd)


def _paragraph_left_border(paragraph, hex_color: str, size: int = 24) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(size))
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), hex_color)
    p_bdr.append(left)
    p_pr.append(p_bdr)


def _set_run_font(run, *, name: str = "Calibri", size: int = 11,
                  color: RGBColor = TEXT, bold: bool = False, italic: bool = False) -> None:
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    for attr in ("ascii", "hAnsi", "cs"):
        rFonts.set(qn(f"w:{attr}"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


# ---------- High-level helpers ----------


def add_cover(doc: Document, title: str, subtitle: str, doc_id: str) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    # Brand strip
    p = doc.add_paragraph()
    _paragraph_shade(p, "1F4E79")
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("  OPEN VALLEY  ")
    _set_run_font(run, name="Segoe UI Semibold", size=14, color=RGBColor(0xFF, 0xFF, 0xFF), bold=True)

    # Document ID strip
    p = doc.add_paragraph()
    _paragraph_shade(p, "2E86AB")
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run(f"  {doc_id}  ")
    _set_run_font(run, name="Segoe UI", size=10, color=RGBColor(0xFF, 0xFF, 0xFF))

    # Title
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(title)
    _set_run_font(run, name="Segoe UI", size=28, color=PRIMARY, bold=True)

    # Subtitle
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(24)
    run = p.add_run(subtitle)
    _set_run_font(run, name="Segoe UI", size=13, color=MUTED, italic=True)


def add_h1(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    _set_run_font(run, name="Segoe UI", size=18, color=PRIMARY, bold=True)
    # underline-style accent bar via bottom border
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "2E86AB")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_h2(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    _set_run_font(run, name="Segoe UI", size=14, color=ACCENT, bold=True)


def add_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    _set_run_font(run, name="Calibri", size=11, color=TEXT)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        _set_run_font(run, name="Calibri", size=11, color=TEXT)


def add_callout(doc: Document, kind: str, body: str) -> None:
    palette = {
        "NOTE": (CALLOUT_NOTE_BG, CALLOUT_NOTE_BAR, "Note"),
        "IMPORTANT": (CALLOUT_IMPORTANT_BG, CALLOUT_IMPORTANT_BAR, "Important"),
        "TIP": (CALLOUT_TIP_BG, CALLOUT_TIP_BAR, "Tip"),
    }
    bg, bar, label = palette[kind]

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(0)
    _paragraph_shade(p, bg)
    _paragraph_left_border(p, bar)
    run = p.add_run(label.upper())
    _set_run_font(run, name="Segoe UI Semibold", size=10,
                  color=RGBColor.from_string(bar), bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    _paragraph_shade(p, bg)
    _paragraph_left_border(p, bar)
    run = p.add_run(body)
    _set_run_font(run, name="Calibri", size=11, color=TEXT)


def add_code_block(doc: Document, code: str, *, lang: str = "") -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    _paragraph_shade(p, CODE_BG)
    _paragraph_left_border(p, "C9D1D9", size=12)
    if lang:
        run = p.add_run(f"{lang}\n")
        _set_run_font(run, name="Consolas", size=8, color=MUTED, italic=True)
    run = p.add_run(code)
    _set_run_font(run, name="Consolas", size=10, color=TEXT)


def add_table(doc: Document, headers: list[str], rows: list[list[str]],
              col_widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False

    # Header
    hdr = table.rows[0].cells
    for idx, head in enumerate(headers):
        cell = hdr[idx]
        _shade(cell, TABLE_HEADER_BG)
        _set_cell_borders(cell, color="1F4E79", size=4)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(head)
        _set_run_font(run, name="Segoe UI Semibold", size=10,
                      color=RGBColor(0xFF, 0xFF, 0xFF), bold=True)

    # Body
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            if r_idx % 2 == 0:
                _shade(cell, ZEBRA_BG)
            _set_cell_borders(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            run = p.add_run(str(value))
            mono = c_idx > 0 and any(ch.isdigit() for ch in str(value)) and "." in str(value)
            font_name = "Consolas" if mono else "Calibri"
            _set_run_font(run, name=font_name, size=10, color=TEXT)

    if col_widths:
        for row in table.rows:
            for c_idx, width in enumerate(col_widths):
                row.cells[c_idx].width = Cm(width)


def add_kv_grid(doc: Document, pairs: list[tuple[str, str]],
                col_widths: tuple[float, float] = (5.5, 11.0)) -> None:
    table = doc.add_table(rows=len(pairs), cols=2)
    table.autofit = False
    for r_idx, (key, value) in enumerate(pairs):
        kc, vc = table.rows[r_idx].cells
        for cell in (kc, vc):
            _set_cell_borders(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _shade(kc, "EEF2F7")

        kc.text = ""
        p = kc.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(key)
        _set_run_font(run, name="Segoe UI Semibold", size=10, color=PRIMARY, bold=True)

        vc.text = ""
        p = vc.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(value)
        _set_run_font(run, name="Calibri", size=11, color=TEXT)

        kc.width = Cm(col_widths[0])
        vc.width = Cm(col_widths[1])


def add_footer(doc: Document, label: str) -> None:
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Open Valley · {label}")
    _set_run_font(run, name="Segoe UI", size=9, color=MUTED, italic=True)


# =============================================================
# Document 1 — Environment Variables and Configuration
# =============================================================


def build_doc_env(out: Path) -> None:
    doc = Document()

    add_cover(
        doc,
        title="Environment Variables & Configuration",
        subtitle="What is needed to run Open Valley locally and to deploy it",
        doc_id="HANDOVER · 01 of 04",
    )

    add_callout(
        doc,
        "NOTE",
        "Open Valley is a 100% client-side UI/UX project. No environment variables are required to "
        "run, build, or deploy the application. The entries below are optional and are only relevant "
        "if the optional Express server in server/ is ever turned on.",
    )

    add_h1(doc, "1. Configuration files in the repository")
    add_table(
        doc,
        headers=["File", "Purpose"],
        rows=[
            [".env", "Local non-secret values. Currently only contains PING_MESSAGE."],
            [".env.example", "Reference template. Copy to .env.local for future backend extensions."],
            [".npmrc", "npm settings. engine-strict=false so installs do not hard-fail on Node minor drift."],
            ["vercel.json", "Vercel build, SPA rewrite, security headers and asset cache headers."],
            ["netlify.toml", "Netlify build command, publish dir and SPA redirect."],
            ["vite.config.ts", "Vite output dir (dist/), manual chunks, terser, esnext target."],
            ["tailwind.config.ts", "Tailwind theme, plugins, content paths."],
            ["postcss.config.js", "PostCSS pipeline (Tailwind + Autoprefixer)."],
            ["tsconfig.json", "TypeScript options and @/ and @shared/ path aliases."],
            ["components.json", "shadcn/ui generator settings."],
        ],
        col_widths=[4.2, 12.0],
    )

    add_h1(doc, "2. Environment variables")
    add_paragraph(doc, "All variables are optional. None must be configured in Vercel or Netlify.")

    add_table(
        doc,
        headers=["Variable", "Required", "Where it is read", "Default"],
        rows=[
            ["PING_MESSAGE", "No", "server/index.ts (GET /api/ping)", '"ping"'],
            ["PORT", "No", "server/node-build.ts (production server)", "3000"],
            ["VITE_API_URL", "No", "Reserved for future backend integration", "(unset)"],
            ["VITE_GOOGLE_ANALYTICS_ID", "No", "Reserved for future analytics", "(unset)"],
            ["VITE_ENABLE_DEMO_MODE", "No", "Reserved feature flag", "(unset)"],
        ],
        col_widths=[4.5, 2.0, 6.5, 3.0],
    )

    add_h2(doc, "Current contents of .env")
    add_code_block(
        doc,
        '# Local environment variables for Open Valley.\n'
        '# Non-secret values only. Do not commit secrets here.\n'
        '# This is a 100% client-side UI/UX project; no env vars are required to run it.\n\n'
        '# Optional: message returned by GET /api/ping (only relevant if running the Express server).\n'
        'PING_MESSAGE="ping pong"\n',
        lang="env",
    )

    add_h1(doc, "3. Deployment platform configuration")
    add_callout(
        doc,
        "TIP",
        "Neither Vercel nor Netlify need any environment variables configured in their dashboards. "
        "Both platforms detect the configuration directly from vercel.json and netlify.toml.",
    )

    add_h2(doc, "Vercel (vercel.json)")
    add_kv_grid(
        doc,
        [
            ("Framework", "Vite (auto-detected)"),
            ("Install command", "npm install --include=dev"),
            ("Build command", "npm run build"),
            ("Output directory", "dist"),
            ("SPA rewrite", "/(.*) -> /index.html"),
            ("Security headers", "X-Content-Type-Options, X-Frame-Options, X-XSS-Protection on all routes"),
            ("Cache headers", "no-cache for index.html; max-age=31536000 immutable for hashed assets"),
            ("Node.js version", "20.x (matches package.json engines)"),
        ],
    )

    add_h2(doc, "Netlify (netlify.toml)")
    add_kv_grid(
        doc,
        [
            ("Build command", "npm run build"),
            ("Publish directory", "dist"),
            ("SPA redirect", "/* -> /index.html (200)"),
        ],
    )

    add_footer(doc, "Environment Variables & Configuration · 01 of 04")
    doc.save(str(out))


# =============================================================
# Document 2 — Technology Stack & Dependencies
# =============================================================


PROD_DEPS_CORE = [
    ("react", "^18.3.1", "18.3.1"),
    ("react-dom", "^18.3.1", "18.3.1"),
    ("react-router-dom", "^6.30.1", "6.30.2"),
    ("@tanstack/react-query", "^5.84.2", "5.90.16"),
    ("react-hook-form", "^7.62.0", "7.69.0"),
    ("@hookform/resolvers", "^5.2.1", "5.2.2"),
    ("framer-motion", "^12.23.12", "12.23.26"),
    ("lucide-react", "^0.539.0", "0.539.0"),
    ("recharts", "^2.12.7", "2.15.4"),
    ("reactflow", "^11.11.4", "11.11.4"),
    ("leaflet", "^1.9.4", "1.9.4"),
    ("react-leaflet", "^4.2.1", "4.2.1"),
    ("elk", "^1.0.0", "1.0.0"),
    ("react-day-picker", "^9.8.1", "9.13.0"),
    ("date-fns", "^4.1.0", "4.1.0"),
    ("embla-carousel-react", "^8.6.0", "8.6.0"),
    ("react-resizable-panels", "^3.0.4", "3.0.6"),
    ("input-otp", "^1.4.2", "1.4.2"),
    ("cmdk", "^1.1.1", "1.1.1"),
    ("sonner", "^1.7.4", "1.7.4"),
    ("vaul", "^1.1.2", "1.1.2"),
    ("next-themes", "^0.4.6", "0.4.6"),
    ("class-variance-authority", "^0.7.1", "0.7.1"),
    ("clsx", "^2.1.1", "2.1.1"),
    ("tailwind-merge", "^2.6.0", "2.6.0"),
    ("tailwindcss-animate", "^1.0.7", "1.0.7"),
    ("uuid", "^13.0.0", "13.0.0"),
    ("xlsx", "^0.18.5", "0.18.5"),
    ("better-sqlite3", "^9.2.2", "9.6.0"),
]

PROD_DEPS_RADIX = [
    ("@radix-ui/react-accordion", "^1.2.11", "1.2.12"),
    ("@radix-ui/react-alert-dialog", "^1.1.14", "1.1.15"),
    ("@radix-ui/react-aspect-ratio", "^1.1.7", "1.1.8"),
    ("@radix-ui/react-avatar", "^1.1.10", "1.1.11"),
    ("@radix-ui/react-checkbox", "^1.3.2", "1.3.3"),
    ("@radix-ui/react-collapsible", "^1.1.11", "1.1.12"),
    ("@radix-ui/react-context-menu", "^2.2.15", "2.2.16"),
    ("@radix-ui/react-dialog", "^1.1.14", "1.1.15"),
    ("@radix-ui/react-dropdown-menu", "^2.1.15", "2.1.16"),
    ("@radix-ui/react-hover-card", "^1.1.14", "1.1.15"),
    ("@radix-ui/react-label", "^2.1.7", "2.1.8"),
    ("@radix-ui/react-menubar", "^1.1.15", "1.1.16"),
    ("@radix-ui/react-navigation-menu", "^1.2.13", "1.2.14"),
    ("@radix-ui/react-popover", "^1.1.14", "1.1.15"),
    ("@radix-ui/react-progress", "^1.1.7", "1.1.8"),
    ("@radix-ui/react-radio-group", "^1.3.7", "1.3.8"),
    ("@radix-ui/react-scroll-area", "^1.2.9", "1.2.10"),
    ("@radix-ui/react-select", "^2.2.5", "2.2.6"),
    ("@radix-ui/react-separator", "^1.1.7", "1.1.8"),
    ("@radix-ui/react-slider", "^1.3.5", "1.3.6"),
    ("@radix-ui/react-slot", "^1.2.3", "1.2.4"),
    ("@radix-ui/react-switch", "^1.2.5", "1.2.6"),
    ("@radix-ui/react-tabs", "^1.1.12", "1.1.13"),
    ("@radix-ui/react-toast", "^1.2.14", "1.2.15"),
    ("@radix-ui/react-toggle", "^1.1.9", "1.1.10"),
    ("@radix-ui/react-toggle-group", "^1.1.10", "1.1.11"),
    ("@radix-ui/react-tooltip", "^1.2.7", "1.2.8"),
]

DEV_DEPS = [
    ("vite", "^7.1.2", "7.3.0"),
    ("@vitejs/plugin-react-swc", "^4.0.0", "4.2.2"),
    ("@swc/core", "^1.13.3", "1.15.8"),
    ("typescript", "^5.9.2", "5.9.3"),
    ("tailwindcss", "^3.4.17", "3.4.19"),
    ("@tailwindcss/typography", "^0.5.16", "0.5.19"),
    ("postcss", "^8.5.6", "8.5.6"),
    ("autoprefixer", "^10.4.21", "10.4.23"),
    ("terser", "^5.36.0", "5.46.0"),
    ("prettier", "^3.6.2", "3.7.4"),
    ("@types/node", "^24.2.1", "24.10.4"),
    ("@types/react", "^18.3.23", "18.3.27"),
    ("@types/react-dom", "^18.3.7", "18.3.7"),
]


def build_doc_stack(out: Path) -> None:
    doc = Document()

    add_cover(
        doc,
        title="Technology Stack & Dependencies",
        subtitle="Languages, frameworks, and every package with its exact installed version",
        doc_id="HANDOVER · 02 of 04",
    )

    add_callout(
        doc,
        "NOTE",
        "Each dependency table shows two version columns: the range declared in package.json and "
        "the exact version that was resolved by package-lock.json on the verified install used to "
        "produce a successful production build.",
    )

    add_h1(doc, "1. Languages and core tooling")
    add_table(
        doc,
        headers=["Layer", "Technology", "Version"],
        rows=[
            ["Language", "TypeScript", "5.9.3"],
            ["UI framework", "React", "18.3.1"],
            ["Build tool", "Vite", "7.3.0"],
            ["Compiler", "@swc/core (via @vitejs/plugin-react-swc)", "1.15.8 / 4.2.2"],
            ["Styling", "Tailwind CSS", "3.4.19"],
            ["CSS processor", "PostCSS", "8.5.6"],
            ["Vendor prefixes", "Autoprefixer", "10.4.23"],
            ["Minifier", "Terser", "5.46.0"],
            ["Formatter", "Prettier", "3.7.4"],
            ["Runtime", "Node.js", "20.x (engines)"],
            ["Package manager", "npm", "10.x"],
        ],
        col_widths=[3.6, 7.5, 5.0],
    )

    add_h1(doc, "2. Runtime dependencies")

    add_h2(doc, "2.1 Core runtime libraries")
    add_table(
        doc,
        headers=["Package", "Range (package.json)", "Installed (lockfile)"],
        rows=[list(t) for t in PROD_DEPS_CORE],
        col_widths=[6.8, 4.8, 4.6],
    )

    add_h2(doc, "2.2 Radix UI primitives")
    add_table(
        doc,
        headers=["Package", "Range (package.json)", "Installed (lockfile)"],
        rows=[list(t) for t in PROD_DEPS_RADIX],
        col_widths=[6.8, 4.8, 4.6],
    )

    add_h1(doc, "3. Development dependencies")
    add_table(
        doc,
        headers=["Package", "Range (package.json)", "Installed (lockfile)"],
        rows=[list(t) for t in DEV_DEPS],
        col_widths=[6.8, 4.8, 4.6],
    )

    add_h1(doc, "4. Lockfile and reproducibility")
    add_paragraph(
        doc,
        "package-lock.json is the single source of truth for exact transitive versions. CI and "
        "Vercel install with npm ci (or npm install --include=dev), which honors the lockfile and "
        "guarantees reproducible builds.",
    )

    add_footer(doc, "Technology Stack & Dependencies · 02 of 04")
    doc.save(str(out))


# =============================================================
# Document 3 — Install / Run / Build / Deploy
# =============================================================


def build_doc_install(out: Path) -> None:
    doc = Document()

    add_cover(
        doc,
        title="Install · Run · Build · Deploy",
        subtitle="Step-by-step instructions for local development and production deployment",
        doc_id="HANDOVER · 03 of 04",
    )

    add_h1(doc, "1. Prerequisites")
    add_table(
        doc,
        headers=["Tool", "Version", "Notes"],
        rows=[
            ["Node.js", "20.x (>= 20.19 recommended)", "Lower 20.x still installs and builds because engine-strict is disabled."],
            ["npm", "10.x", "Bundled with Node 20."],
            ["Git", "any recent version", "For cloning and version control."],
        ],
        col_widths=[3.0, 4.5, 9.0],
    )

    add_h1(doc, "2. Local development")
    add_paragraph(doc, "Install dependencies and start the Vite dev server with hot reloading:")
    add_code_block(
        doc,
        "npm install\n"
        "npm run dev\n",
        lang="bash",
    )
    add_callout(doc, "TIP", "The dev server listens on http://localhost:8080 (configured in vite.config.ts).")

    add_h1(doc, "3. Production build")
    add_paragraph(doc, "Produce optimized static assets in dist/, then optionally preview them locally:")
    add_code_block(
        doc,
        "npm run build\n"
        "npm run preview   # optional\n",
        lang="bash",
    )
    add_h2(doc, "Build outputs")
    add_table(
        doc,
        headers=["Output", "Path"],
        rows=[
            ["Entry HTML", "dist/index.html"],
            ["JS bundles", "dist/assets/*.js (chunked: index, recharts, radix-ui)"],
            ["CSS bundle", "dist/assets/*.css"],
            ["Static assets", "dist/favicon.ico, dist/open-valley-logo.png, dist/placeholder.svg, dist/robots.txt"],
        ],
        col_widths=[4.5, 11.5],
    )

    add_h1(doc, "4. Deploy to Vercel (recommended)")
    add_paragraph(
        doc,
        "The repo includes a complete vercel.json. Vercel will auto-detect the configuration, so "
        "no settings need to be entered manually.",
    )
    add_kv_grid(
        doc,
        [
            ("Framework preset", "Vite"),
            ("Install command", "npm install --include=dev"),
            ("Build command", "npm run build"),
            ("Output directory", "dist"),
            ("Node.js version", "20.x"),
        ],
    )
    add_h2(doc, "Steps")
    add_bullets(doc, [
        "Push the repository to GitHub (current default branch: main).",
        "In the Vercel dashboard, click Add New > Project and import the repo.",
        "Accept the auto-detected settings (sourced from vercel.json).",
        "Click Deploy. Subsequent pushes to main deploy automatically.",
    ])
    add_callout(
        doc,
        "IMPORTANT",
        "vercel.json already provides the SPA rewrite (so React Router works on hard refresh), "
        "security headers (X-Content-Type-Options, X-Frame-Options, X-XSS-Protection), long-cache "
        "headers for hashed assets, and a no-cache header for index.html so deploys take effect "
        "immediately.",
    )

    add_h1(doc, "5. Deploy to Netlify (alternative)")
    add_paragraph(doc, "The repo includes netlify.toml with the correct build command and SPA redirect:")
    add_code_block(
        doc,
        '[build]\n'
        '  command = "npm run build"\n'
        '  publish = "dist"\n\n'
        '[[redirects]]\n'
        '  force = true\n'
        '  from = "/*"\n'
        '  to = "/index.html"\n'
        '  status = 200\n',
        lang="toml",
    )

    add_h1(doc, "6. Deploy to DigitalOcean App Platform (alternative)")
    add_kv_grid(
        doc,
        [
            ("Build command", "npm run build"),
            ("Output directory", "dist"),
        ],
    )

    add_h1(doc, "7. Available npm scripts")
    add_table(
        doc,
        headers=["Script", "Purpose"],
        rows=[
            ["npm run dev", "Start the Vite dev server with HMR."],
            ["npm run build", "Produce production assets in dist/."],
            ["npm run preview", "Serve the production build locally."],
        ],
        col_widths=[4.5, 11.5],
    )

    add_h1(doc, "8. Project layout")
    add_code_block(
        doc,
        "Open-Valley/\n"
        "├── client/              # Application source (React + TypeScript)\n"
        "│   ├── App.tsx          # Router, providers, route table\n"
        "│   ├── main.tsx         # React mount point\n"
        "│   ├── pages/           # Route-level pages\n"
        "│   ├── components/      # Reusable UI + feature components\n"
        "│   │   ├── ui/          # shadcn/ui + Radix-based primitives\n"
        "│   │   ├── access-control/\n"
        "│   │   ├── audit/\n"
        "│   │   └── reports/\n"
        "│   ├── contexts/        # React contexts\n"
        "│   ├── hooks/           # Custom hooks\n"
        "│   ├── i18n/            # Locale files\n"
        "│   ├── lib/             # Utilities\n"
        "│   ├── services/        # Mock data services\n"
        "│   └── global.css       # Tailwind layer + global tokens\n"
        "├── public/              # Static assets copied into dist/\n"
        "├── server/              # Optional Express server (not used in static deploy)\n"
        "├── shared/              # Types shared between client and server\n"
        "├── netlify/functions/   # Optional serverless wrapper for the server\n"
        "├── .project-rules/      # Project rule files\n"
        "├── index.html           # Vite HTML entry\n"
        "├── vite.config.ts       # Vite configuration\n"
        "├── tailwind.config.ts   # Tailwind configuration\n"
        "├── tsconfig.json        # TypeScript configuration\n"
        "├── vercel.json          # Vercel deployment config\n"
        "├── netlify.toml         # Netlify deployment config\n"
        "└── package.json\n",
    )

    add_h1(doc, "9. Build verification checklist")
    add_bullets(doc, [
        "npm install --include=dev completes without errors (warnings only).",
        "npm run build produces dist/index.html and dist/assets/*.",
        "Open dist/index.html via npm run preview and verify the dashboard loads.",
        "Click through major routes (Dashboard, Analytics, Alarms, Automation, Topology, Reports, "
        "Access Control, Settings) and verify navigation and browser refresh both work.",
    ])

    add_footer(doc, "Install · Run · Build · Deploy · 03 of 04")
    doc.save(str(out))


# =============================================================
# Document 4 — Third-Party Services & Integrations
# =============================================================


def build_doc_third_party(out: Path) -> None:
    doc = Document()

    add_cover(
        doc,
        title="Third-Party Services & Integrations",
        subtitle="External services, APIs, and credentials used by the deployed application",
        doc_id="HANDOVER · 04 of 04",
    )

    add_callout(
        doc,
        "IMPORTANT",
        "None. Open Valley is a UI/UX-only project. There are no API keys, no authentication "
        "providers, no analytics services, no payment processors, and no managed databases used by "
        "the deployed application. Therefore there is nothing to transfer during handover for this "
        "section: no accounts to migrate, no API tokens to rotate, no third-party billing to assume.",
    )

    add_h1(doc, "1. Verified absent")
    add_table(
        doc,
        headers=["Category", "Used?", "Notes"],
        rows=[
            ["Authentication providers (Auth0, Clerk, Firebase Auth, Supabase Auth, etc.)",
             "No",
             "Login screen uses local mock auth via client/contexts/AuthContext.tsx."],
            ["Analytics (GA, Mixpanel, PostHog, Segment, etc.)",
             "No",
             "No tracking scripts; reserved env var only."],
            ["Error tracking (Sentry, Rollbar, etc.)",
             "No",
             "Not wired."],
            ["Maps / tiles (Mapbox, Google Maps)",
             "No",
             "Uses leaflet + react-leaflet with default OpenStreetMap tiles only."],
            ["External REST/GraphQL APIs",
             "No",
             "All data rendered in the app is mock data bundled into the client."],
            ["Payment / billing",
             "No",
             "Not applicable."],
            ["Email / SMS / push services",
             "No",
             "Not wired."],
            ["CDN / asset hosts",
             "No",
             "Static assets served by the deployment platform (Vercel / Netlify)."],
            ["Database",
             "No",
             "better-sqlite3 is declared but unused at runtime; the deployed bundle is purely static."],
        ],
        col_widths=[5.5, 2.0, 8.5],
    )

    add_h1(doc, "2. Network calls made by the deployed site")
    add_paragraph(
        doc,
        "The only network calls a user's browser will make on the deployed site are:",
    )
    add_bullets(doc, [
        "Loading the static assets from the deployment domain (Vercel or Netlify CDN).",
        "OpenStreetMap tile requests issued by leaflet when a map view is opened. These are free, "
        "unauthenticated, and require no API key.",
    ])
    add_callout(
        doc,
        "TIP",
        "If the tile provider is ever swapped (e.g., Mapbox), an API key would then become required. "
        "Document any future provider in this section: provider name, what data is exchanged, where "
        "the credential lives (env var name + deployment platform secret), the owner of the account, "
        "and the rotation/transfer procedure.",
    )

    add_h1(doc, "3. Handover checklist")
    add_bullets(doc, [
        "No accounts to transfer (no auth/analytics/error tracking providers).",
        "No API keys to rotate.",
        "No third-party billing to assume.",
        "No webhook endpoints to update.",
        "Source-code repository (GitHub) and the chosen deployment platform (Vercel or Netlify) are "
        "the only external accounts associated with this project.",
    ])

    add_footer(doc, "Third-Party Services & Integrations · 04 of 04")
    doc.save(str(out))


# =============================================================
# Entry point
# =============================================================


def main() -> None:
    root = Path(__file__).resolve().parent.parent
    out_dir = root / "handover-docs"
    out_dir.mkdir(exist_ok=True)

    targets = [
        ("01-Environment-Variables-and-Configuration.docx", build_doc_env),
        ("02-Technology-Stack-and-Dependencies.docx", build_doc_stack),
        ("03-Install-Run-Build-Deploy.docx", build_doc_install),
        ("04-Third-Party-Services-and-Integrations.docx", build_doc_third_party),
    ]
    for filename, builder in targets:
        path = out_dir / filename
        builder(path)
        print(f"Wrote {path}")


if __name__ == "__main__":
    main()
